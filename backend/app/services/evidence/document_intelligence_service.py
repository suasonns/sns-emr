"""AI Document Intelligence for uploaded clinical documents (UCIER extension).

Uploaded documents (labs, H&Ps, hospital records, imaging reports, etc.)
are a distinct evidence source from authored notes -- the text isn't
written by SNS staff, it has to be *extracted* from a file first. This
module is the extraction + classification layer that feeds into the same
`harvest_from_source` orchestrator used for every other source.

Design contract (same as ai_extraction_service.py):
    - NEVER raises. Any failure (unsupported file type, extraction error,
      AI call failure) is logged and degrades gracefully -- the uploaded
      file and DocumentRecord row are always preserved regardless.
    - NEVER fabricates. Only surfaces findings explicitly present in the
      document text/image.
    - Inert/no-op if Azure OpenAI is not configured (reuses the same
      AZURE_OPENAI_* env vars as ai_extraction_service.py).

Extraction strategy:
    - text/plain            -> decode directly.
    - application/pdf        -> pypdf text-layer extraction, page by page.
      Any individual pages with no extractable text (scanned/faxed pages
      mixed into an otherwise text-based chart export) are OCR'd
      individually via Azure Document Intelligence (prebuilt-read) when
      AZURE_DOCINTEL_* is configured, and spliced back into the right
      position -- this is NOT all-or-nothing at the whole-document level.
    - .docx                  -> python-docx paragraph text.
    - image/jpeg|png|tiff    -> no local OCR step; the raw image bytes are
      sent directly to the AI as a vision input, since the same call also
      does classification/extraction (single round trip).
    - If OCR is not configured (or the OCR call itself fails) and a page has
      no text layer, that page contributes no text and the whole document is
      flagged `needs_manual_review=True` so staff know extraction may be
      incomplete -- it is never silently dropped without a signal.
"""

from __future__ import annotations

import base64
import json
import logging
import os
import time
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any

import httpx

logger = logging.getLogger("sns_emr")

DEFAULT_TIMEOUT_SECONDS = 45.0
MAX_SOURCE_TEXT_CHARS = 12000

# Azure Document Intelligence (OCR) -- used only as a fallback for PDFs with
# no extractable text layer (scanned/image-only documents).
AZURE_DOCINTEL_API_VERSION = "2024-11-30"
OCR_POLL_INTERVAL_SECONDS = 2.0
OCR_POLL_MAX_ATTEMPTS = 30  # ~60s max wait for OCR to finish

IMAGE_CONTENT_TYPES = {"image/jpeg", "image/png", "image/tiff"}
TEXT_LAYER_CONTENT_TYPES = {"application/pdf"}
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PLAIN_TEXT_CONTENT_TYPE = "text/plain"

# Keep in sync with document_storage.MIME_TYPE_EXTENSIONS -- these are the
# only content types the upload endpoint accepts in the first place.
SUPPORTED_CONTENT_TYPES = (
    IMAGE_CONTENT_TYPES
    | TEXT_LAYER_CONTENT_TYPES
    | {DOCX_CONTENT_TYPE, PLAIN_TEXT_CONTENT_TYPE, "application/msword"}
)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    method: str  # "text_layer" | "ocr" | "docx" | "plain_text" | "vision" | "unsupported"
    needs_manual_review: bool = False
    image_base64: str | None = None  # populated only when method == "vision"


def extract_text_from_file(
    *, file_bytes: bytes, content_type: str, file_name: str | None = None
) -> ExtractionResult:
    """Best-effort text extraction from an uploaded document's raw bytes.

    Never raises -- any parsing failure degrades to needs_manual_review.
    """

    try:
        if content_type == PLAIN_TEXT_CONTENT_TYPE:
            text = file_bytes.decode("utf-8", errors="replace").strip()
            return ExtractionResult(text=text, method="plain_text")

        if content_type in TEXT_LAYER_CONTENT_TYPES:
            return _extract_pdf_text(file_bytes)

        if content_type == DOCX_CONTENT_TYPE:
            return _extract_docx_text(file_bytes)

        if content_type in IMAGE_CONTENT_TYPES:
            return ExtractionResult(
                text="",
                method="vision",
                image_base64=base64.b64encode(file_bytes).decode("ascii"),
            )

        # application/msword (legacy .doc) and anything else -- no pure-python
        # extractor available; requires manual review for now.
        logger.info(
            "document_intelligence: unsupported content_type=%s for text extraction "
            "(file_name=%s) -- flagged for manual review",
            content_type,
            file_name,
        )
        return ExtractionResult(text="", method="unsupported", needs_manual_review=True)
    except Exception:
        logger.exception(
            "document_intelligence: extraction failed file_name=%s content_type=%s",
            file_name,
            content_type,
        )
        return ExtractionResult(text="", method="unsupported", needs_manual_review=True)


def _extract_pdf_text(file_bytes: bytes) -> ExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    page_texts: list[str] = []
    empty_page_indices: list[int] = []
    for i, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        page_texts.append(page_text)
        if not page_text:
            empty_page_indices.append(i)

    if not empty_page_indices:
        return ExtractionResult(text="\n\n".join(page_texts).strip(), method="text_layer")

    # One or more pages have no extractable text layer -- these are
    # scanned/faxed/image-only pages mixed into an otherwise-text PDF (very
    # common for hospital chart exports). Rather than treating the whole
    # document as one all-or-nothing case, OCR just the affected pages via
    # Azure Document Intelligence and splice the results back in at the
    # right position, so nothing silently goes missing.
    ocr_page_texts = _ocr_pdf_pages_via_document_intelligence(file_bytes, empty_page_indices)
    ocr_used = False
    if ocr_page_texts:
        ocr_used = True
        for idx, text in ocr_page_texts.items():
            if text:
                page_texts[idx] = text

    joined = "\n\n".join(t for t in page_texts if t).strip()
    still_missing = [i for i in empty_page_indices if not page_texts[i]]

    if not joined:
        # Nothing usable at all (no text layer anywhere, and OCR unavailable
        # or produced nothing) -- flag for manual review rather than guess.
        return ExtractionResult(text="", method="text_layer", needs_manual_review=True)

    return ExtractionResult(
        text=joined,
        method="ocr" if ocr_used else "text_layer",
        # Even with a partial OCR success, flag for review if any page
        # still has no text (e.g. OCR failed on some or was unavailable) --
        # staff should know the extracted text may be incomplete.
        needs_manual_review=bool(still_missing),
    )


def _azure_docintel_config() -> dict[str, str] | None:
    endpoint = os.getenv("AZURE_DOCINTEL_ENDPOINT")
    api_key = os.getenv("AZURE_DOCINTEL_KEY")

    if not (endpoint and api_key):
        return None

    return {"endpoint": endpoint.rstrip("/"), "api_key": api_key}


def is_ocr_configured() -> bool:
    return _azure_docintel_config() is not None


def _ocr_pdf_pages_via_document_intelligence(
    file_bytes: bytes, page_indices: list[int]
) -> dict[int, str] | None:
    """OCR only the given 0-based page indices of a PDF via Azure Document
    Intelligence's prebuilt-read model, returning {original_page_index: text}.

    Builds a small sub-PDF containing just those pages (preserving order) so
    a single OCR call covers every affected page, then maps the per-page
    results in the response back to the original page indices using each
    page's character spans into the response's full `content` string.

    Never raises -- returns None if OCR is not configured, or on any
    submit/poll/parse/PDF-assembly failure.
    """

    config = _azure_docintel_config()
    if config is None:
        logger.info("document_intelligence: OCR fallback skipped (Azure Document Intelligence not configured)")
        return None

    try:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(BytesIO(file_bytes))
        writer = PdfWriter()
        for idx in page_indices:
            writer.add_page(reader.pages[idx])

        sub_pdf_buffer = BytesIO()
        writer.write(sub_pdf_buffer)
        sub_pdf_bytes = sub_pdf_buffer.getvalue()
    except Exception:
        logger.exception("document_intelligence: failed to assemble sub-PDF of empty-text pages for OCR")
        return None

    analyze_url = (
        f"{config['endpoint']}/documentintelligence/documentModels/prebuilt-read:analyze"
        f"?api-version={AZURE_DOCINTEL_API_VERSION}"
    )

    try:
        submit = httpx.post(
            analyze_url,
            headers={
                "Ocp-Apim-Subscription-Key": config["api_key"],
                "Content-Type": "application/pdf",
            },
            content=sub_pdf_bytes,
            timeout=DEFAULT_TIMEOUT_SECONDS,
        )
        submit.raise_for_status()
        operation_url = submit.headers.get("operation-location")
        if not operation_url:
            logger.warning("document_intelligence: OCR submit response missing operation-location header")
            return None

        analyze_result: dict[str, Any] | None = None
        for _ in range(OCR_POLL_MAX_ATTEMPTS):
            time.sleep(OCR_POLL_INTERVAL_SECONDS)
            poll = httpx.get(
                operation_url,
                headers={"Ocp-Apim-Subscription-Key": config["api_key"]},
                timeout=DEFAULT_TIMEOUT_SECONDS,
            )
            poll.raise_for_status()
            body = poll.json()
            status = body.get("status")
            if status == "succeeded":
                analyze_result = body.get("analyzeResult") or {}
                break
            if status == "failed":
                logger.warning("document_intelligence: OCR analysis failed: %r", body.get("error"))
                return None
        else:
            logger.warning("document_intelligence: OCR polling timed out waiting for result")
            return None

        if analyze_result is None:
            return None

        content = analyze_result.get("content") or ""
        ocr_pages = analyze_result.get("pages") or []

        result: dict[int, str] = {}
        for ocr_page in ocr_pages:
            sub_pdf_page_number = ocr_page.get("pageNumber")  # 1-based, in sub-PDF order
            if not sub_pdf_page_number or sub_pdf_page_number < 1 or sub_pdf_page_number > len(page_indices):
                continue
            original_index = page_indices[sub_pdf_page_number - 1]

            spans = ocr_page.get("spans") or []
            page_text_parts = [content[s["offset"] : s["offset"] + s["length"]] for s in spans if "offset" in s]
            page_text = "".join(page_text_parts).strip()
            if page_text:
                result[original_index] = page_text

        return result or None
    except Exception:
        logger.exception("document_intelligence: OCR call failed")
        return None


def _extract_docx_text(file_bytes: bytes) -> ExtractionResult:
    import docx

    document = docx.Document(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Also pull table cell text (labs/vitals are frequently tabular in DOCX).
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    joined = "\n".join(paragraphs).strip()
    if not joined:
        return ExtractionResult(text="", method="docx", needs_manual_review=True)
    return ExtractionResult(text=joined, method="docx")


@dataclass(frozen=True)
class KeyFinding:
    label: str
    value: str
    category: str | None = None  # e.g. "lab_result", "diagnosis", "functional_status"
    original_text_excerpt: str = ""


@dataclass(frozen=True)
class DocumentIntelligenceResult:
    document_type_guess: str
    summary: str
    key_findings: list[KeyFinding] = field(default_factory=list)
    confidence: float | None = None


def _azure_openai_config() -> dict[str, str] | None:
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")

    if not (endpoint and api_key and api_version and deployment):
        return None

    return {
        "endpoint": endpoint.rstrip("/"),
        "api_key": api_key,
        "api_version": api_version,
        "deployment": deployment,
    }


def is_configured() -> bool:
    return _azure_openai_config() is not None


_SYSTEM_PROMPT = """You are a clinical document intake reviewer for a hospice agency's \
EMR. Hospice staff upload external documents (labs, History & Physical notes, hospital \
discharge summaries, imaging reports, insurance/authorization letters, advance \
directives, physician orders, etc.) and you help staff quickly understand what each \
document is and what it says, WITHOUT ever diagnosing or interpreting beyond what is \
explicitly written.

Non-negotiable rules:
- You NEVER diagnose, interpret lab results as good/bad/normal/abnormal beyond what the \
  document itself states, or infer a patient's prognosis.
- You NEVER fabricate a value, date, or finding that is not present in the document.
- Every key finding you output must be directly supported by a short verbatim excerpt \
  from the document.
- If you cannot confidently identify the document type, use "OTHER" and explain why in \
  the summary.
- If the document contains no clinically meaningful content (e.g. a blank cover page, \
  a fax header only), return an empty "key_findings" array.

Document type MUST be one of: "H_AND_P", "LABS_DIAGNOSTICS", "DISCHARGE_SUMMARY", \
"HOSPITAL_RECORDS", "IMAGING_REPORT", "PHYSICIAN_ORDERS", "INSURANCE_AUTHORIZATION", \
"ADVANCE_DIRECTIVE", "CONSULT_NOTE", "OTHER".

For each key finding, use category one of: "lab_result", "diagnosis", "functional_status", \
"decline_indicator", "medication", "vital_sign", "imaging_finding", "administrative", \
"other". For lab results, the "value" field should include the result and units/flag \
exactly as documented (e.g. "Sodium 128 mEq/L (LOW)").

Respond ONLY with a JSON object of this exact shape:
{
  "document_type_guess": "one of the enum values above",
  "summary": "1-3 sentence plain-language summary of what this document is and contains.",
  "confidence": 0.0-1.0,
  "key_findings": [
    {
      "label": "short label, e.g. 'Sodium' or 'Primary diagnosis'",
      "value": "the documented value/finding",
      "category": "one of the category enum values above",
      "original_text_excerpt": "verbatim short excerpt (<= 300 chars) supporting this"
    }
  ]
}
"""


def classify_and_extract_document(
    *,
    text: str | None,
    image_base64: str | None = None,
    content_type: str,
    hint_document_type: str | None = None,
) -> DocumentIntelligenceResult | None:
    """Classify an uploaded document and extract its key findings.

    Returns None if unconfigured, on any API/parse error, or if there is
    no usable content (empty text and no image) -- never raises.
    """

    config = _azure_openai_config()
    if config is None:
        logger.info("document_intelligence: AI classification skipped (Azure OpenAI not configured)")
        return None

    cleaned_text = (text or "").strip()
    if not cleaned_text and not image_base64:
        return None

    user_context_prefix = f"uploaded_document_type_hint: {hint_document_type or 'unknown'}\n"

    if image_base64:
        user_content: Any = [
            {"type": "text", "text": user_context_prefix + "Review the attached document image."},
            {
                "type": "image_url",
                "image_url": {"url": f"data:{content_type};base64,{image_base64}"},
            },
        ]
    else:
        truncated_text = cleaned_text[:MAX_SOURCE_TEXT_CHARS]
        user_content = user_context_prefix + f"--- DOCUMENT TEXT ---\n{truncated_text}"

    url = (
        f"{config['endpoint']}/openai/deployments/{config['deployment']}"
        f"/chat/completions?api-version={config['api_version']}"
    )

    payload = {
        "messages": [
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
    }

    try:
        response = httpx.post(
            url,
            headers={"api-key": config["api_key"], "Content-Type": "application/json"},
            json=payload,
            timeout=DEFAULT_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        body = response.json()
        raw_content = body["choices"][0]["message"]["content"]
        parsed = json.loads(raw_content)
    except Exception:
        logger.exception("document_intelligence: AI classification call failed")
        return None

    return _parse_result(parsed)


def _parse_result(parsed: Any) -> DocumentIntelligenceResult | None:
    if not isinstance(parsed, dict):
        return None

    document_type_guess = str(parsed.get("document_type_guess") or "OTHER").strip().upper()
    summary = str(parsed.get("summary") or "").strip()

    confidence_raw = parsed.get("confidence")
    confidence: float | None
    try:
        confidence = float(confidence_raw) if confidence_raw is not None else None
        if confidence is not None:
            confidence = max(0.0, min(1.0, confidence))
    except (TypeError, ValueError):
        confidence = None

    findings_raw = parsed.get("key_findings")
    findings: list[KeyFinding] = []
    if isinstance(findings_raw, list):
        for item in findings_raw:
            try:
                finding = _parse_finding(item)
            except Exception:
                logger.warning("document_intelligence: skipping malformed key_finding item=%r", item)
                continue
            if finding is not None:
                findings.append(finding)

    return DocumentIntelligenceResult(
        document_type_guess=document_type_guess,
        summary=summary,
        key_findings=findings,
        confidence=confidence,
    )


def _parse_finding(item: Any) -> KeyFinding | None:
    if not isinstance(item, dict):
        return None

    label = str(item.get("label") or "").strip()
    value = str(item.get("value") or "").strip()
    if not label or not value:
        return None

    return KeyFinding(
        label=label[:128],
        value=value[:1000],
        category=(str(item.get("category"))[:32] if item.get("category") else None),
        original_text_excerpt=str(item.get("original_text_excerpt") or "").strip()[:2000],
    )
